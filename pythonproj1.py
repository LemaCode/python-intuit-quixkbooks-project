from docx import Document
from docx.shared import Inches
import pyttsx3

# run: pip3 install -r requirments.txt
# to install the require package to use thos progrom

# text to speeck function
def speak(text):
   pyttsx3.speak(text)


document = Document()

# profile picture
document.add_picture('BJP Chapter 8.png', width=Inches(22.0))

# identity
name = input('What is your name? ')
speak('Hello' + name + ' how are your today?')
speak('What is your phone number? ')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')
document.add_paragraph(
   name + ' | ' + phone_number + ' | ' + email)

# about me
document.add_heading('About me')
about_me = input('Tell about yourself? ')
document.add_paragraph(about_me)

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company')
from_date = input('From Date')
to_date = input('To Date')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
   'Describe your experience at ' + company)
p.add_run(experience_details)

# skills
document.add_heading('Skills')

skill = input('Enter skills')
sk = document.add_paragraph(skill)
sk.style = 'List Bullet'

# more skills
while True:
   has_more_skills = input('Do you have more skills? Yes or No')

   if has_more_skills.lower() == 'yes':
   
      skill = input('Enter skills')
      sk = document.add_paragraph(skill)
      sk.style = 'List Bullet'
   else:
      break

# more experiences
while True:
   has_more_experiences = input(
      'Do you have more experiences? Yes or No')

   if has_more_experiences.lower() == 'yes':
      p = document.add_paragraph()
   
      company = input('Enter company')
      from_date = input('From Date')
      to_date = input('To Date')
   
      p.add_run(company + ' ').bold = True
      p.add_run(from_date + '-' + to_date + '\n').italic = True
   
      experience_details = input(
         'Describe your experience at ' + company)
      p.add_run(experience_details)
   
   else:
      break

        # footer
section = document.sections[0]
footer = section.footer
f = footer.paragraph[0]
f.text = "CV generator project"
document.save('cv.docx')